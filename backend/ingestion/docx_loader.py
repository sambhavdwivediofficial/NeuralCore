# ingestion/docx_loader.py
from __future__ import annotations

import asyncio
import io
from typing import Any

from ingestion.base_loader import BaseLoader, SourceType
from ingestion.loader_factory import register_loader


@register_loader(SourceType.DOCX)
class DocxLoader(BaseLoader):
    source_type = SourceType.DOCX

    async def load(self, source_config: dict[str, Any]) -> list[dict[str, Any]]:
        raw_bytes = await self._read_bytes(source_config)
        text, extra_metadata = await asyncio.to_thread(self._parse_docx, raw_bytes)
        source_id = source_config.get("file_path")
        metadata = {"source_type": self.source_type.value, "file_path": source_id, **extra_metadata}
        return [self._build_document(text, metadata=metadata, source_id=source_id)]

    @staticmethod
    def _parse_docx(raw_bytes: bytes) -> tuple[str, dict[str, Any]]:
        from docx import Document

        document = Document(io.BytesIO(raw_bytes))
        parts: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style else ""
            if style_name.startswith("Heading"):
                level = "".join(filter(str.isdigit, style_name)) or "1"
                parts.append(f"{'#' * min(int(level), 6)} {text}")
            else:
                parts.append(text)

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))

        core_props = document.core_properties
        metadata = {
            "title": core_props.title or None,
            "author": core_props.author or None,
            "created": core_props.created.isoformat() if core_props.created else None,
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
        }
        return "\n\n".join(parts), metadata
